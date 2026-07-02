import argparse
import json
import os
from collections import defaultdict
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# Загружаем JSON файл
def load_json(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError("Файл не найден: " + file_path)
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict) or not data:
        raise ValueError("JSON-файл пустой или имеет неверную структуру")
    return data

# Ищем файл с данными
def resolve_input_file(preferred="result_plan.json"):
    candidates = [preferred, "result_plan(1).json", "result_plan.json"]
    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate
    return preferred

# Убираем лишние пробелы и переносы строк
def normalize_text(value):
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\r", " ").replace("\n", " ")
    while "  " in text:
        text = text.replace("  ", " ")
    return text.strip()

# Получаем количество кредитов
def get_credits(discipline):
    for key in ("credits", "total_credits"):
        try:
            value = discipline.get(key, 0)
            if value is not None:
                return float(value)
        except (TypeError, ValueError):
            continue
    return 0.0

# Загружаем соответствие код -> номер модуля
def load_module_config():
    candidates = ["module_config(1).json", "module_config.json"]
    module_map = {}
    for file_path in candidates:
        if not os.path.exists(file_path):
            continue
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            disciplines = data.get("disciplines", {})
            for item in disciplines.values():
                code = normalize_text(item.get("code", ""))
                module = item.get("module")
                if code and module is not None:
                    module_map[code] = module
            return module_map
        except:
            continue
    return {}

# Загружаем названия модулей
def load_module_names():
    candidates = ["module_config(1).json", "module_config.json"]
    module_names = {}
    for file_path in candidates:
        if not os.path.exists(file_path):
            continue
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            modules = data.get("modules", {})
            for key, value in modules.items():
                try:
                    module_num = int(key)
                    module_names[module_num] = value.get("name", "Модуль " + str(module_num))
                except:
                    continue
            return module_names
        except:
            continue
    return {}

# Делаем перенос длинных слов
def wrap_text(text, max_length=25):
    if len(text) <= max_length:
        return text
    words = text.split()
    lines = []
    current_line = []
    current_length = 0
    for word in words:
        if len(word) > max_length:
            if current_line:
                lines.append(" ".join(current_line))
                current_line = []
                current_length = 0
            for i in range(0, len(word), max_length - 5):
                chunk = word[i:i + max_length - 5]
                if chunk:
                    lines.append(chunk + "-")
            continue
        if current_length + len(word) + (1 if current_line else 0) <= max_length:
            current_line.append(word)
            current_length = current_length + len(word) + (1 if current_line else 0)
        else:
            if current_line:
                lines.append(" ".join(current_line))
            current_line = [word]
            current_length = len(word)
    if current_line:
        lines.append(" ".join(current_line))
    return "\n".join(lines)

# Подготавливаем семестры: убираем факультативы и null, сортируем
def prepare_semestrs(semestrs, module_map):
    prepared = {}
    for sem_key, disciplines in semestrs.items():
        filtered = []
        for disc in disciplines:
            if disc.get("is_facultative", False):
                continue
            disc = dict(disc)
            code = normalize_text(disc.get("code", ""))
            module = module_map.get(code)
            if module is None:
                continue
            disc["module"] = module
            filtered.append(disc)
        def sort_key(item):
            return (item.get("module"), -get_credits(item), normalize_text(item.get("name", "")).lower())
        filtered.sort(key=sort_key)
        prepared[str(sem_key)] = filtered
    return prepared

# Группируем семестры по курсам
def group_by_course(semestrs):
    grouped = defaultdict(lambda: defaultdict(list))
    for sem_str, disciplines in semestrs.items():
        try:
            sem = int(sem_str)
        except:
            continue
        course = (sem + 1) // 2
        for disc in disciplines:
            grouped[course][sem].append(disc)
    return dict(grouped)

# Ищем максимальное количество кредитов
def find_max_credits(semestrs):
    max_credits = 0.0
    for disciplines in semestrs.values():
        for disc in disciplines:
            credits = get_credits(disc)
            if credits > max_credits:
                max_credits = credits
    return max_credits if max_credits > 0 else 1.0

# Склоняем слово "год"
def plural_years(value):
    text = normalize_text(value)
    if not text:
        return "не указан"
    try:
        number = int(float(text))
        if number % 10 == 1 and number % 100 != 11:
            suffix = "год"
        elif number % 10 in (2, 3, 4) and number % 100 not in (12, 13, 14):
            suffix = "года"
        else:
            suffix = "лет"
        return str(number) + " " + suffix
    except:
        return text

# Заливка ячейки цветом
def add_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

# Отступы внутри ячейки
def set_cell_margins(cell, top=60, start=60, bottom=60, end=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    sides = [("top", top), ("start", start), ("bottom", bottom), ("end", end)]
    for side, value in sides:
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

# Высота строки
def set_row_height(row, height_pt):
    row.height = Pt(height_pt)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

# Поля страницы
def set_doc_margins(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(0.7)
    section.right_margin = Cm(0.7)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)

# Настройка шрифта
def set_font(run, size=11, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

# Добавление абзаца с текстом
def add_text_paragraph(container, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, underline=False):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size, bold)
    run.font.underline = underline
    return p

# Очистка ячейки
def clear_cell(cell):
    cell.text = ""

# Заполнение ячейки текстом
def fill_cell_text(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color_fill=None, module=None):
    clear_cell(cell)
    
    # Разделяем текст на части (название и кредиты)
    parts = text.split("\n")
    name = parts[0] if len(parts) > 0 else ""
    credits_text = parts[1] if len(parts) > 1 else ""
    
    # Основной абзац с текстом
    p = cell.paragraphs[0]
    p.alignment = align  # Используем переданное выравнивание
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    
    # Если текст нужно выровнять по центру, не делаем перенос
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        # Для заголовков просто выводим текст без переноса
        run = p.add_run(text)
        set_font(run, size, bold)
    else:
        # Для дисциплин делаем умный перенос
        wrapped_name = wrap_text(name, 25)
        name_lines = wrapped_name.split("\n")
        for i, line in enumerate(name_lines):
            if i > 0:
                p.add_run().add_break()
            run = p.add_run(line)
            set_font(run, size, bold)
        
        # Кредиты - отдельный абзац (выравнивание по правому краю)
        if credits_text:
            credits_p = cell.add_paragraph()
            credits_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            credits_p.paragraph_format.space_before = Pt(0)
            credits_p.paragraph_format.space_after = Pt(0)
            credits_p.paragraph_format.line_spacing = 1.0
            credits_run = credits_p.add_run(credits_text)
            set_font(credits_run, size, bold)
    
    # Номер модуля (только для дисциплин, не для заголовков)
    if module is not None and align != WD_ALIGN_PARAGRAPH.CENTER:
        module_p = cell.add_paragraph()
        module_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        module_p.paragraph_format.space_before = Pt(10)
        module_p.paragraph_format.space_after = Pt(0)
        module_p.paragraph_format.line_spacing = 1.0
        module_run = module_p.add_run(str(module))
        set_font(module_run, 6, False)
    
    # Выравнивание по вертикали
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if align == WD_ALIGN_PARAGRAPH.CENTER else WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    
    if color_fill:
        add_shading(cell, color_fill)

# Стиль таблицы
def add_table_border_style(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

# Создание документа
def build_document(data, output_path):
    semestrs_raw = data.get("semestrs", {})
    if not semestrs_raw:
        raise ValueError("В JSON нет данных о семестрах")
    
    # Загружаем модули
    module_map = load_module_config()
    module_names = load_module_names()
    
    # Подготавливаем семестры
    semestrs = prepare_semestrs(semestrs_raw, module_map)
    
    # Определяем номера семестров
    all_semesters = []
    for s in semestrs.keys():
        if str(s).isdigit():
            all_semesters.append(int(s))
    all_semesters.sort()
    
    if not all_semesters:
        raise ValueError("Не удалось определить номера семестров")
    
    total_semesters = len(all_semesters)
    max_course = 0
    for sem in all_semesters:
        course = (sem + 1) // 2
        if course > max_course:
            max_course = course
    
    max_credits = find_max_credits(semestrs)
    
    # Считаем максимальное число строк
    max_rows = 0
    for disciplines in semestrs.values():
        if len(disciplines) > max_rows:
            max_rows = len(disciplines)
    
    if max_rows <= 0:
        max_rows = 1
    
    # Создаём документ
    doc = Document()
    section = doc.sections[0]
    set_doc_margins(section)
    
    # Базовый стиль
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(10)
    
    # Шапка
    add_text_paragraph(
        doc,
        "ВИЗУАЛИЗАЦИЯ УЧЕБНОГО ПЛАНА",
        14,
        True,
        WD_ALIGN_PARAGRAPH.CENTER,
        True
    )
    doc.add_paragraph()
    
    # Информация о плане
    plan_info = data.get("plan_info", {})
    title = normalize_text(plan_info.get("title", "Не указано"))
    qualification = normalize_text(plan_info.get("qualification", "Не указана"))
    duration = plural_years(plan_info.get("duration", ""))
    year = normalize_text(plan_info.get("year", "Не указан"))
    
    add_text_paragraph(doc, "Направление: " + (title or "не указано"), 10, True)
    add_text_paragraph(doc, "Квалификация: " + (qualification or "не указана"), 10, True)
    add_text_paragraph(doc, "Срок обучения: " + duration, 10, True)
    add_text_paragraph(doc, "Учебный год: " + (year or "не указан"), 10, True)
    
    doc.add_paragraph()
    
    # Создаём таблицу
    table = doc.add_table(rows=2 + max_rows, cols=total_semesters)
    add_table_border_style(table)
    
    usable_width_cm = 29.7 - 0.7 - 0.7
    col_width_cm = usable_width_cm / total_semesters
    for col in range(total_semesters):
        table.columns[col].width = Cm(col_width_cm)
    
    # Заголовки курсов и семестров
    for col_index, sem in enumerate(all_semesters):
        course = (sem + 1) // 2
        
        # Первая строка - название курса (объединяем ячейки)
        top_cell = table.cell(0, col_index)
        if sem % 2 == 1 and col_index + 1 < total_semesters:
            top_cell = top_cell.merge(table.cell(0, col_index + 1))
        
        # Для заголовка курса используем выравнивание по центру
        fill_cell_text(
            top_cell,
            "КУРС " + str(course),
            9,
            True,
            WD_ALIGN_PARAGRAPH.CENTER,  # Выравнивание по центру
            "D9D9D9"
        )
        
        # Вторая строка - номер семестра
        bottom_cell = table.cell(1, col_index)
        fill_cell_text(
            bottom_cell,
            "Сем. " + str(sem),
            9,
            True,
            WD_ALIGN_PARAGRAPH.CENTER,  # Выравнивание по центру
            "DDEBF7"
        )
    
    # Заполняем таблицу дисциплинами
    base_height_per_credit = 12.0
    minimum_height = 20.0
    
    for row_index in range(max_rows):
        row_credits = 0.0
        for col_index, sem in enumerate(all_semesters):
            disciplines = semestrs.get(str(sem), [])
            cell = table.cell(row_index + 2, col_index)
            
            if row_index < len(disciplines):
                disc = disciplines[row_index]
                name = normalize_text(disc.get("name", ""))
                credits = get_credits(disc)
                module = disc.get("module")
                
                if credits > row_credits:
                    row_credits = credits
                
                # Форматируем кредиты
                if credits == int(credits):
                    credits_str = str(int(credits))
                else:
                    credits_str = str(credits)
                
                text = name + "\n(" + credits_str + " зет)"
                fill_cell_text(
                    cell,
                    text,
                    9,
                    False,
                    WD_ALIGN_PARAGRAPH.LEFT,  # Для дисциплин - по левому краю
                    None,
                    module
                )
            else:
                fill_cell_text(cell, "", 9, False, WD_ALIGN_PARAGRAPH.LEFT)
        
        # Высота строки
        if max_credits > 0:
            row_height = max(minimum_height, base_height_per_credit * (row_credits / max_credits))
        else:
            row_height = minimum_height
        set_row_height(table.rows[row_index + 2], row_height)
    
    # Высота заголовков
    set_row_height(table.rows[0], 18)
    set_row_height(table.rows[1], 16)
    
    # Список модулей
    doc.add_paragraph()
    add_text_paragraph(
        doc,
        "СПИСОК МОДУЛЕЙ:",
        11,
        True,
        WD_ALIGN_PARAGRAPH.LEFT,
        True
    )
    doc.add_paragraph()
    
    for module_num in sorted(module_names.keys()):
        module_name = module_names.get(module_num, "Модуль " + str(module_num))
        add_text_paragraph(
            doc,
            str(module_num) + ". " + module_name,
            10,
            False,
            WD_ALIGN_PARAGRAPH.LEFT
        )
    
    doc.save(output_path)
    print("Готово: " + output_path)

# Запуск программы
def main():
    parser = argparse.ArgumentParser(description="Визуализация учебного плана в DOCX")
    parser.add_argument("-i", "--input", default=None, help="Путь к JSON-файлу")
    parser.add_argument("-o", "--output", default="visualized_plan.docx", help="Путь к DOCX")
    args = parser.parse_args()
    input_file = args.input
    if input_file is None:
        input_file = resolve_input_file()
    try:
        data = load_json(input_file)
        build_document(data, args.output)
    except FileNotFoundError as e:
        print("Ошибка: " + str(e))
    except ValueError as e:
        print("Ошибка данных: " + str(e))
    except Exception as e:
        print("Неожиданная ошибка: " + str(e))

if __name__ == "__main__":
    main()